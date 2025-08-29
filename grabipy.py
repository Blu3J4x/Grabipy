#!/usr/bin/env python3
"""
Current Version: 1.5
Grabipy – A robust and user-friendly Python script for threat intelligence.

This tool automatically scans files and folders to extract Indicators of Compromise (IOCs)
including IPs, hashes (MD5, SHA1, SHA256), domains, URLs, and email addresses.
"""
import json
import os, re, ipaddress, requests, csv, time, getpass, sys, subprocess, email, struct
from urllib.parse import urlparse, unquote
from collections import Counter, defaultdict
from datetime import datetime, timedelta
import base64
import hashlib
import configparser
from requests.exceptions import RequestException

# === Cache Handling ===
CACHE_FILE = 'enrichment_cache.json'
CACHE_EXPIRY_HOURS = 24

def load_cache():
    """Loads the enrichment cache from a JSON file if it exists."""
    if not os.path.exists(CACHE_FILE):
        return {}
    try:
        with open(CACHE_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (json.JSONDecodeError, IOError):
        # If the file is corrupted or unreadable, start with an empty cache
        return {}

def save_cache(cache_data):
    """Saves the enrichment cache to a JSON file."""
    try:
        with open(CACHE_FILE, 'w', encoding='utf-8') as f:
            json.dump(cache_data, f, indent=4)
    except IOError as e:
        tqdm.write(f"{color.ERROR}[✗] Failed to save enrichment cache: {e}{color.END}")

# === Console colors ===
class color:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKCYAN = '\033[96m'
    OKGREEN = '\033[92m'  # Add this line
    SUCCESS = '\033[92m' # Add this line too for consistency
    WARNING = '\033[93m'
    ERROR = '\033[91m'
    END = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'
    INFO = '\033[94m'

# === Dependency check ===
def check_and_install_dependencies():
    """Checks for all required dependencies and installs them in a single batch."""
    
    # Import standard libraries needed for this function
    import sys
    import subprocess

    # A single dictionary mapping the import name to the pip install name
    all_dependencies = {
        "requests": "requests",      # For API calls
        "chardet": "chardet",        # For file encoding detection
        "docx": "python-docx",
        "pdfplumber": "pdfplumber",
        "extract_msg": "extract-msg",
        "pandas": "pandas",
        "openpyxl": "openpyxl",
        "tqdm": "tqdm",
        "tldextract": "tldextract",
        "bs4": "beautifulsoup4",
        "scapy": "scapy",
        "six": "six",
        "numpy": "numpy",
        "dateutil": "python-dateutil",
        "pytz": "pytz"
    }
    
    missing_packages = []
    print(f"{color.INFO}[*] Checking for required modules...{color.END}")
    for module_name, pip_name in all_dependencies.items():
        try:
            __import__(module_name)
        except ImportError:
            missing_packages.append(pip_name)
            
    if missing_packages:
        print(f"\n{color.WARNING}[!] The following required modules are missing:{color.END}")
        for pkg in missing_packages:
            print(f"    - {pkg}")
            
        choice = input(f"\nDo you want to try installing them now? [Y/n]: ").strip().lower()
        if choice in ('', 'y', 'yes'):
            try:
                print(f"{color.INFO}[*] Installing {len(missing_packages)} package(s)...{color.END}")
                # Installs all missing packages in one command
                subprocess.check_call([sys.executable, "-m", "pip", "install", "--user", *missing_packages])
                print(f"\n{color.OKGREEN}[✓] All dependencies installed. Please re-run the script to start.{color.END}")
            except subprocess.CalledProcessError:
                print(f"{color.ERROR}[✗] Failed to install one or more packages. Please try installing them manually.{color.END}")
            sys.exit(1) # Exit after installation attempt so the script can be re-run cleanly
        else:
            print(f"{color.ERROR}[✗] Required modules are missing. Exiting.{color.END}")
            sys.exit(1)
    else:
        print(f"{color.OKGREEN}[✓] All dependencies are satisfied.{color.END}")

# NOTE: The dependency check is now called inside the `if __name__ == "__main__"` block

# Safe imports
from docx import Document
import olefile
import chardet
import pdfplumber
import extract_msg
import pandas as pd
from tqdm import tqdm
import tldextract
from bs4 import BeautifulSoup
from scapy.all import rdpcap, IP, TCP, UDP, Raw, hexdump
from scapy.layers.http import HTTP

# === Config & API Keys ===
CONFIG_FILE = 'config.ini'

def setup_config():
    """Prompts user for API keys and saves them to a config file."""
    config = configparser.ConfigParser()
    config['API_KEYS'] = {}
    print(f"\n{color.INFO}[*] Please enter your API keys. They will be saved to '{CONFIG_FILE}'{color.END}")
    config['API_KEYS']['AbuseIPDB_Key'] = getpass.getpass("Enter AbuseIPDB API key: ").strip()
    config['API_KEYS']['VirusTotal_Key'] = getpass.getpass("Enter VirusTotal API key: ").strip()

    with open(CONFIG_FILE, 'w') as configfile:
        config.write(configfile)
    print(f"{color.SUCCESS}[✓] API keys saved to '{CONFIG_FILE}'. You will not be asked for them again.{color.END}")
    return config['API_KEYS']['AbuseIPDB_Key'], config['API_KEYS']['VirusTotal_Key']

def load_config():
    """Loads API keys from a config file."""
    config = configparser.ConfigParser()
    if not os.path.exists(CONFIG_FILE):
        return None, None
    config.read(CONFIG_FILE)
    try:
        abuse_key = config['API_KEYS']['AbuseIPDB_Key']
        vt_key = config['API_KEYS']['VirusTotal_Key']
        return abuse_key, vt_key
    except (KeyError, configparser.NoSectionError):
        print(f"{color.ERROR}[✗] Error reading API keys from '{CONFIG_FILE}'. Please re-enter them.{color.END}")
        return None, None

# === API Endpoints ===
ABUSEIPDB_URL = 'https://api.abuseipdb.com/api/v2/check'
VT_FILE_URL = 'https://www.virustotal.com/api/v3/files/'
VT_DOMAIN_URL = 'https://www.virustotal.com/api/v3/domains/'
VT_URLS_URL = 'https://www.virustotal.com/api/v3/urls/'

# === Regex (More specific hash regex for improved accuracy) ===
RE_IPV4 = re.compile(r"(?:[0-9]{1,3}\.){3}[0-9]{1,3}")
RE_MD5 = re.compile(r"\b[a-fA-F0-9]{32}\b")
RE_SHA1 = re.compile(r"\b[a-fA-F0-9]{40}\b")
RE_SHA256 = re.compile(r"\b[a-fA-F0-9]{64}\b")
RE_HASH = re.compile(f"(?:{RE_MD5.pattern}|{RE_SHA1.pattern}|{RE_SHA256.pattern})")
RE_EMAIL = re.compile(r"\b[a-zA-Z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
RE_URL = re.compile(r"\b(?:http|https|hxxp|hxxps)://\S+\b", re.IGNORECASE)

# NEW: AbuseIPDB Category Map
cat_map = {
    3: "Fraud Orders", 4: "DDoS Attack", 5: "FTP Brute-Force", 6: "Ping of Death",
    7: "Phishing", 8: "Fraud VoIP", 9: "Open Proxy", 10: "Web Spam",
    11: "Email Spam", 12: "Blog Spam", 13: "VPN IP", 14: "Port Scan",
    15: "Hacking", 16: "SQL Injection", 17: "Spoofing", 18: "Brute-Force",
    19: "Bad Web Bot", 20: "Exploited Host", 21: "Web App Attack",
    22: "SSH", 23: "IoT Targeted"
}

# === Risk scoring ===
def get_risk_level(score):
    if score >= 70: return "High"
    elif score >= 30: return "Medium"
    else: return "Low"

# === Helpers ===
def undefang(url: str) -> str:
    url = url.replace("hxxp://", "http://").replace("hxxps://", "https://")
    url = url.replace("[.]", ".").replace("(.)", ".")
    url = url.replace("2F", "/").replace("40", "@").replace("%/", "/").replace("%/%/", "//")
    return url
    
def defang(ioc_type: str, ioc: str) -> str:
    """Defangs an IOC to prevent it from being clickable."""
    if ioc_type == 'IP':
        return ioc.replace(".", "[.]")
    elif ioc_type == 'Domain':
        return ioc.replace(".", "[.]")
    elif ioc_type == 'URL':
        url = ioc.replace("http://", "hxxp://").replace("https://", "hxxps://")
        url = url.replace(".", "[.]")
        return url
    return ioc

def get_root_domain(url_or_email):
    url_or_email = url_or_email.strip().strip("<>").strip('"')
    if "@" not in url_or_email:
        parsed = tldextract.extract(undefang(url_or_email))
        if parsed.domain and parsed.suffix: return f"{parsed.domain}.{parsed.suffix}"
        return url_or_email
    else:
        domain = url_or_email.split("@")[-1]
        parsed = tldextract.extract(domain)
        if parsed.domain and parsed.suffix: return f"{parsed.domain}.{parsed.suffix}"
        return domain

def clean_email(email_address):
    """Extracts the email address from a string."""
    if not email_address: return ""
    match = re.search(RE_EMAIL, email_address)
    return match.group(0) if match else ""
    
def extract_safelink(url):
    """Extracts the real URL from a Microsoft SafeLink wrapper."""
    if 'safelinks.protection.outlook.com' in url.lower():
        try:
            parsed = urlparse(url)
            query_params = dict(qp.split('=') for qp in parsed.query.split('&') if '=' in qp)
            if 'url' in query_params:
                unwrapped_url = unquote(query_params['url'])
                return unwrapped_url
        except Exception as e:
            tqdm.write(f"{color.WARNING}[!] Failed to parse SafeLink URL: {url} - {e}{color.END}")
    return url
    
def clean_message_id(message_id):
    """Removes the leading and trailing angle brackets from a Message-ID."""
    if message_id and message_id.startswith('<') and message_id.endswith('>'):
        return message_id[1:-1]
    return message_id

# === Generic IOC check ===
def is_generic_ioc(ioc, ioc_type):
    """Checks if an IOC belongs to a known generic/trusted entity."""
    
    # Common domains for Google, Microsoft, AWS, etc.
    GENERIC_DOMAINS = {
        "google.com", "gmail.com", "gstatic.com", "googlevideo.com", 
        "googleusercontent.com", "youtube.com", "windows.net", "microsoft.com", 
        "msftncsi.com", "azure.com", "office.com", "outlook.com", "hotmail.com", 
        "live.co.uk", "amazonaws.com", "s3.amazonaws.com", "cloudfront.net", 
        "akamaihd.net", "apple.com", "icloud.com", "github.com", "mailchimp.com", 
        "dropbox.com", "onedrive.live.com", "logmein.com", "facebook.com", 
        "twitter.com", "linkedin.com"
    }
    
    # Common IP ranges for Google, Microsoft, AWS, etc.
    GENERIC_IP_RANGES = [
        # Google
        "8.8.8.0/24", "8.8.4.0/24", "172.217.0.0/16", "142.250.0.0/16", "216.58.192.0/19",
        # Microsoft (Azure/Office 365)
        "40.76.0.0/16", "52.96.0.0/11", "13.107.6.152/31", "13.107.18.152/31",
        # Amazon AWS
        "3.0.0.0/5", "18.200.0.0/13", "34.192.0.0/12", "52.0.0.0/8", "54.240.0.0/12",
        "72.0.0.0/8", "76.0.0.0/8", "99.80.0.0/12", "107.20.0.0/14", "204.236.192.0/18",
        # Cloudflare
        "104.16.0.0/12", "172.64.0.0/13", "188.114.96.0/20",
    ]

    # Hashes are never generic, so return False
    if ioc_type == "Hash":
        return False
        
    # Check IPs against CIDR blocks
    if ioc_type == "IP":
        try:
            ip_obj = ipaddress.ip_address(ioc)
            for network in GENERIC_IP_RANGES:
                if ip_obj in ipaddress.ip_network(network):
                    return True
        except ValueError:
            pass # Invalid IP, not generic
            
    # Check Domains and URLs against known domains
    if ioc_type in ["Domain", "URL"]:
        root_domain = get_root_domain(ioc)
        if root_domain in GENERIC_DOMAINS:
            return True

    return False

# === File collector ===
SUPPORTED_EXTENSIONS = ['.txt', '.csv', '.xlsx', '.xls', '.docx', '.pdf', '.msg', '.eml', '.pcap']

def get_files_from_path(path):
    path = os.path.expanduser(path)
    path = os.path.normpath(path)
    if os.path.isfile(path):
        ext = os.path.splitext(path)[1].lower()
        if ext in SUPPORTED_EXTENSIONS: return [path]
        else: raise ValueError(f"Unsupported file type: {path}")
    elif os.path.isdir(path):
        all_files = []
        for root, dirs, files in os.walk(path):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in SUPPORTED_EXTENSIONS:
                    all_files.append(os.path.join(root, file))
        return all_files
    else:
        raise ValueError(f"Path does not exist: {path}")

# === File hashing function ===
def hash_file(filepath):
    """Calculates MD5, SHA1, and SHA256 hashes for a given file."""
    md5_hash = hashlib.md5()
    sha1_hash = hashlib.sha1()
    sha256_hash = hashlib.sha256()
    
    try:
        with open(filepath, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                md5_hash.update(chunk)
                sha1_hash.update(chunk)
                sha256_hash.update(chunk)
        
        return [
            {'hash': md5_hash.hexdigest(), 'type': 'md5', 'source_file': filepath},
            {'hash': sha1_hash.hexdigest(), 'type': 'sha1', 'source_file': filepath},
            {'hash': sha256_hash.hexdigest(), 'type': 'sha256', 'source_file': filepath}
        ]
    except Exception as e:
        tqdm.write(f"{color.ERROR}[!] Error hashing file {filepath}: {e}{color.END}")
        return []

# === File parsing ===
def read_file_content(file_path):
    """Generator to read file content line by line for memory efficiency."""
    try:
        if file_path.lower().endswith((".txt", ".csv")):
            # --- Auto-detect encoding ---
            with open(file_path, 'rb') as f_raw:
                raw_data = f_raw.read(1024) # Read a sample for detection
            detected = chardet.detect(raw_data)
            encoding = detected.get('encoding', 'utf-8')
            tqdm.write(f"{color.INFO}[*] Detected encoding for {os.path.basename(file_path)}: {encoding}{color.END}")

            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                for line in f:
                    yield line.strip()

        # ... (rest of the function for other file types remains the same) ...
        elif pd and file_path.lower().endswith(('.xlsx', '.xls')):
            excel_file = pd.ExcelFile(file_path)
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)
                for row in df.itertuples(index=False):
                    yield ' '.join([str(cell) for cell in row if pd.notna(cell)])
        elif Document and file_path.lower().endswith('.docx'):
            doc = Document(file_path)
            for para in doc.paragraphs:
                yield para.text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield cell.text
        elif pdfplumber and file_path.lower().endswith('.pdf'):
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        for line in text.splitlines():
                            yield line

        elif olefile and extract_msg and file_path.lower().endswith(('.msg', '.eml')):
            yield f"Email file placeholder: {file_path}"

        elif rdpcap and file_path.lower().endswith('.pcap'):
            yield f"PCAP file placeholder: {file_path}"
            
    except Exception as e:
        tqdm.write(f"{color.ERROR}[✗] Error reading {file_path}: {e}{color.END}")

def read_eml_file(file_path, scan_attachments=False):
    """Specialized function for parsing .eml files."""
    text, attachments, headers_info, html_urls = [], [], [], []
    try:
        import email.policy
        with open(file_path, "rb") as f:
            msg = email.message_from_binary_file(f, policy=email.policy.default)
        headers_info.append({
            'From': msg.get('From'),
            'Reply-To': msg.get('Reply-To'),
            'Return-Path': msg.get('Return-Path'),
            'Message-ID': msg.get('Message-ID'),
            'X-Sender-IP': msg.get('X-Sender-IP')
        })
        
        def decode_payload(part):
            payload = part.get_payload(decode=True)
            if not payload: return "", ""
            charset = part.get_content_charset() or "utf-8"
            try: 
                return payload.decode(charset, errors="ignore"), part.get_content_type()
            except: 
                return payload.decode("utf-8", errors="ignore"), part.get_content_type()
        
        for part in msg.walk():
            decoded_content, ctype = decode_payload(part)
            disp = part.get_content_disposition()
            
            if disp == "attachment" and scan_attachments:
                payload = part.get_payload(decode=True)
                filename = part.get_filename() or "unknown_filename"
                if payload: attachments.append((payload, filename))
            
            if decoded_content:
                if ctype == "text/html":
                    soup = BeautifulSoup(decoded_content, 'html.parser')
                    for a_tag in soup.find_all('a', href=True):
                        html_urls.append(a_tag['href'])
                    for img_tag in soup.find_all('img', src=True):
                        html_urls.append(img_tag['src'])
                    text.extend(soup.get_text().splitlines())
                elif ctype == "text/plain":
                    text.extend(decoded_content.splitlines())
                    
    except Exception as e:
        tqdm.write(f"{color.ERROR}[✗] Error reading {file_path}: {e}{color.END}")
        
    return ([line.strip() for line in text if line and line.strip()], 
            attachments, 
            headers_info, 
            html_urls)

def read_pcap_file_and_extract_files(file_path):
    """
    Reads a PCAP file and extracts IOCs (IPs, Domains, URLs) and also attempts to
    reconstruct and save transferred files.
    NOTE: This only works for unencrypted HTTP traffic.
    """
    ips = set()
    domains = set()
    urls = set()
    extracted_files = []

    try:
        packets = rdpcap(file_path)
        sessions = packets.sessions()
        
        for session_key, session_packets in sessions.items():
            if 'TCP' in session_key:
                src_ip, src_port, dst_ip, dst_port = session_key.split()
                if dst_port == '80':
                    data = b''
                    for pkt in session_packets:
                        if pkt.haslayer(Raw):
                            data += pkt[Raw].load

                    if b'HTTP/1.' in data:
                        try:
                            http_stream = HTTP(data)
                            if http_stream.haslayer('HTTP Response') and hasattr(http_stream, 'Content-Type'):
                                content_type = http_stream.Content_Type.decode('utf-8', errors='ignore')
                                content_length = int(http_stream.Content_Length.decode('utf-8')) if hasattr(http_stream, 'Content-Length') else 0
                                payload = data.split(b'\r\n\r\n', 1)[1]
                                
                                if payload and content_length > 0 and len(payload) >= content_length:
                                    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_extracted_files")
                                    os.makedirs(output_dir, exist_ok=True)
                                    filename = f"extracted_{dst_ip}_{dst_port}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
                                    if 'image/jpeg' in content_type: filename += '.jpg'
                                    elif 'image/png' in content_type: filename += '.png'
                                    elif 'application/pdf' in content_type: filename += '.pdf'
                                    elif 'application/zip' in content_type: filename += '.zip'
                                    elif 'text/plain' in content_type: filename += '.txt'
                                    else: filename += '.bin'
                                    
                                    output_path = os.path.join(output_dir, filename)
                                    with open(output_path, 'wb') as f:
                                        f.write(payload)
                                    extracted_files.append(output_path)
                                    tqdm.write(f"{color.SUCCESS}[✓] Extracted file: {output_path}{color.END}")

                        except Exception as e:
                            tqdm.write(f"{color.WARNING}[!] Failed to process HTTP stream for {session_key}: {e}{color.END}")

            for pkt in session_packets:
                if pkt.haslayer(IP):
                    ips.add(pkt[IP].src)
                    ips.add(pkt[IP].dst)

                if pkt.haslayer('DNS'):
                    for qname in pkt['DNS'].qd:
                        qname_str = qname.qname.decode().strip('.')
                        ext = tldextract.extract(qname_str)
                        if ext.domain and ext.suffix:
                            domains.add(f"{ext.domain}.{ext.suffix}")
                
                if pkt.haslayer(TCP) and (pkt[TCP].dport == 80 or pkt[TCP].sport == 80):
                    if pkt.haslayer(Raw):
                        try:
                            http_payload = pkt[Raw].load.decode('utf-8', errors='ignore')
                            match = re.search(r"Host: (.*?)\r\n", http_payload)
                            if match:
                                host = match.group(1).strip()
                                urls.add(f"http://{host}")
                        except UnicodeDecodeError:
                            pass
        
        return list(ips), list(domains), list(urls), extracted_files

    except Exception as e:
        tqdm.write(f"{color.ERROR}[✗] Error reading {file_path}: {e}{color.END}")
        return [], [], [], []

# === IOC Extraction ===
def extract_iocs_from_file(file_path, scan_attachments):
    ips = set()
    hashes = set()
    domains = set()
    urls = set()
    emails_dict = {}
    msg_ids = set()
    extracted_pcap_files = []
    
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == ".pcap":
        pcap_ips, pcap_domains, pcap_urls, extracted_pcap_files = read_pcap_file_and_extract_files(file_path)
        ips.update(pcap_ips)
        domains.update(pcap_domains)
        urls.update(pcap_urls)
        for h_dict in [h for f in extracted_pcap_files for h in hash_file(f)]:
            hashes.add((h_dict['hash'], h_dict['type'], h_dict['source_file']))
    else:
        if file_extension in (".eml", ".msg"):
            lines, attachments, headers_info, html_urls = read_eml_file(file_path, scan_attachments)
        else:
            lines_generator = read_file_content(file_path)
            lines = [l for l in lines_generator]
            attachments, headers_info, html_urls = [], [], []

        def add_url_and_domain(u):
            u_defanged = undefang(u)
            u_final = extract_safelink(u_defanged)
            urls.add(u_final)
            parsed = urlparse(u_final)
            if parsed.hostname:
                add_domain_from_string(parsed.hostname)

        def add_domain_from_string(domain_string):
            domains.add(get_root_domain(domain_string))

        for u in html_urls:
            add_url_and_domain(u)

        for header in headers_info:
            for field in ['From','Reply-To','Return-Path']:
                addr = clean_email(header.get(field))
                if addr:
                    emails_dict[addr] = header.get('Message-ID','')
                    add_domain_from_string(addr.split('@')[-1])
            if header.get('X-Sender-IP'): ips.add(header['X-Sender-IP'])
            if header.get('Message-ID'): msg_ids.add(header['Message-ID'])

        for line in lines:
            line = line.split('#',1)[0].strip()
            if not line: continue
            item = undefang(line)

# --- NEW, MORE ROBUST PARSING LOGIC ---

            # Find all potential IOCs on the line using a lenient regex
            potential_ips = RE_IPV4.findall(item)
            # Validate each potential IP before adding it to the list
            for pot_ip in potential_ips:
                try:
                    ipaddress.ip_address(pot_ip)
                    ips.add(pot_ip)
                except ValueError:
                    # This was not a valid IP, so we ignore it
                    pass
            hashes.update([(h.lower(), 'md5', file_path) for h in RE_MD5.findall(item)])
            hashes.update([(h.lower(), 'sha1', file_path) for h in RE_SHA1.findall(item)])
            hashes.update([(h.lower(), 'sha256', file_path) for h in RE_SHA256.findall(item)])

            for e in RE_EMAIL.findall(item):
                ce = clean_email(e)
                emails_dict[ce] = ''
                add_domain_from_string(ce.split('@')[-1])

            for u in RE_URL.findall(item):
                add_url_and_domain(u)

            # Split the line into words and check each word individually.
            # This is great for finding domains and IPv6 addresses.
            words = re.split(r'[\s,;<>\[\]\(\)]+', item)
            for word in words:
                if not word:
                    continue
                # Check if the word is a valid IP (handles IPv6)
                try:
                    ipaddress.ip_address(word)
                    ips.add(word)
                    continue # Move to the next word
                except ValueError:
                    pass
                
                # Check if the word is a domain
                if '.' in word and '@' not in word:
                    extracted = tldextract.extract(word)
                    if extracted.domain and extracted.suffix:
                        add_domain_from_string(word)
        
        for att_bytes, att_name in attachments:
            source_name = f"{file_path} > {att_name}"
            hashes.add((hashlib.md5(att_bytes).hexdigest(), 'md5', source_name))
            hashes.add((hashlib.sha1(att_bytes).hexdigest(), 'sha1', source_name))
            hashes.add((hashlib.sha256(att_bytes).hexdigest(), 'sha256', source_name))
        
        # <<<FIXED>>>: Removed the unconditional hashing of source files from this function.
        # Hashing is now handled ONLY in the main extract_iocs function, based on user choice.

    return (list(ips), 
            [{'hash': h[0], 'type': h[1], 'source_file': h[2]} for h in hashes], 
            list(domains), 
            list(urls), 
            emails_dict, 
            list(msg_ids),
            extracted_pcap_files)
            
# === Enricher class ===
class Enricher:
    def __init__(self, abuse_key, vt_key, cache, abuse_delay=1.5, vt_delay=16):
        self.abuse_key, self.vt_key = abuse_key, vt_key
        self.cache = cache
        self.abuse_delay, self.vt_delay = abuse_delay, vt_delay
        self.last_abuse_call = 0
        self.last_vt_call = 0
        self.ssl_error_detected = False

    def _get_from_cache(self, ioc):
        """Checks for a valid, non-expired IOC in the cache."""
        if ioc in self.cache:
            cached_item = self.cache[ioc]
            timestamp = datetime.fromisoformat(cached_item['timestamp'])
            if datetime.now() - timestamp < timedelta(hours=CACHE_EXPIRY_HOURS):
                return cached_item['data'], True
        return None, False

    def _update_cache(self, ioc, data):
        """Updates the cache with new enrichment data."""
        self.cache[ioc] = {
            'timestamp': datetime.now().isoformat(),
            'data': data
        }

    def _handle_ssl_error(self):
        """Prints a detailed, one-time warning message about SSL errors."""
        if not self.ssl_error_detected:
            tqdm.write(f"\n{color.ERROR}[✗] CRITICAL SSL ERROR: Certificate verification failed.{color.END}")
            tqdm.write(f"{color.WARNING}   This is common on corporate networks with SSL inspection proxies.{color.END}")
            tqdm.write(f"{color.WARNING}   To fix this, you may need to modify the script to use 'verify=False' in the 'Enricher' class,")
            tqdm.write(f"{color.WARNING}   or provide a path to your company's root certificate.{color.END}")
            tqdm.write(f"{color.INFO}[*] Halting all further enrichment attempts.{color.END}\n")
        self.ssl_error_detected = True
        return {'Error': 'SSL verification failed.'}

    def _wait_for_api(self, api_name):
        if api_name == 'AbuseIPDB':
            delay = self.abuse_delay
            last_call = self.last_abuse_call
        else:
            delay = self.vt_delay
            last_call = self.last_vt_call
        
        elapsed = time.time() - last_call
        if elapsed < delay:
            time.sleep(delay - elapsed)

        if api_name == 'AbuseIPDB':
            self.last_abuse_call = time.time()
        else:
            self.last_vt_call = time.time()

    def enrich_ip(self, ip):
        if self.ssl_error_detected: return {'Error': 'SSL verification failed; skipping further enrichment.'}
        cached_data, was_cached = self._get_from_cache(ip)
        if was_cached: return cached_data, True
        
        self._wait_for_api('AbuseIPDB')
        try:
            check_resp = requests.get(ABUSEIPDB_URL, headers={'Accept':'application/json','Key':self.abuse_key}, params={'ipAddress':ip}, timeout=30)
            check_resp.raise_for_status()
            data = check_resp.json()['data']
            score = data.get('abuseConfidenceScore', 0)
            
            result = {
                'Abuse Score': score,
                'Risk Level': get_risk_level(score),
                'Country': data.get('countryCode',''),
                'ISP': data.get('isp',''),
                'Domain': data.get('domain',''),
                'Hostname(s)': ", ".join(data.get('hostnames',[])) if data.get('hostnames') else '',
                'Last Reported': data.get('lastReportedAt','')
            }

            if score > 25:
                self._wait_for_api('AbuseIPDB')
                reports_url = 'https://api.abuseipdb.com/api/v2/reports'
                reports_resp = requests.get(reports_url, headers={'Accept':'application/json','Key':self.abuse_key}, params={'ipAddress':ip, 'maxAgeInDays': '90'}, timeout=30)
                
                if reports_resp.status_code == 200:
                    reports_data = reports_resp.json().get('data', {}).get('reports', [])
                    
                    if reports_data:
                        category_counts = Counter(cat for report in reports_data for cat in report['categories'])
                        categories_summary = ", ".join([f"{cat_map.get(cat, 'Unknown')} ({count})" for cat, count in category_counts.most_common(3)])
                        result['Report Categories'] = categories_summary

            self._update_cache(ip, result)
            return result, False
            
        except requests.exceptions.SSLError: return self._handle_ssl_error(), False
        except RequestException as e: return {'Error':str(e)}, False
    
    def enrich_hash(self, value):
        if self.ssl_error_detected: return {'Error': 'SSL verification failed; skipping further enrichment.'}
        cached_data, was_cached = self._get_from_cache(value)
        if was_cached: return cached_data, True

        self._wait_for_api('VirusTotal')
        try:
            headers = {"x-apikey": self.vt_key}
            resp = requests.get(VT_FILE_URL + value, headers=headers, timeout=30)
            
            # Raise an HTTPError for bad responses (4xx or 5xx)
            resp.raise_for_status()

            j = resp.json().get('data', {}).get('attributes', {})
            stats = j.get('last_analysis_stats', {})
            malicious_count = int(stats.get('malicious', 0))
            note = ", ".join(j.get('tags', [])) if 'tags' in j else ''
            
            result = {
                'Detection Ratio': f"{malicious_count}/{sum(stats.values())}" if stats else '0/0',
                'Harmless': stats.get('harmless', 0),
                'Malicious': stats.get('malicious', 0),
                'Suspicious': stats.get('suspicious', 0),
                'Undetected': stats.get('undetected', 0),
                'Type Description': j.get('type_description', ''),
                'First Seen': j.get('first_submission_date', ''),
                'Last Seen': j.get('last_submission_date', ''),
                'Tags': note
            }
            self._update_cache(value, result)
            return result, False
            
        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 404:
                # Modified: Create a result for "Not Found" and cache it
                tqdm.write(f"{color.WARNING}[!] Hash {value} not found on VirusTotal. Caching this result.{color.END}")
                result = {'Detection Ratio': 'Not Found', 'Malicious': 'N/A', 'Tags': 'Not Found on VT'}
                self._update_cache(value, result)
                return result, False
            else:
                tqdm.write(f"{color.ERROR}[✗] HTTP error for hash {value}: {e}{color.END}")
                return {'Error': str(e)}, False
        except requests.exceptions.SSLError:
            return self._handle_ssl_error(), False
        except RequestException as e:
            tqdm.write(f"{color.ERROR}[✗] An unexpected error occurred for hash {value}: {e}{color.END}")
            return {'Error': str(e)}, False

    def enrich_domain(self, value):
        if self.ssl_error_detected: return {'Error': 'SSL verification failed; skipping further enrichment.'}
        cached_data, was_cached = self._get_from_cache(value)
        if was_cached: return cached_data, True

        self._wait_for_api('VirusTotal')
        try:
            headers = {"x-apikey": self.vt_key}
            resp = requests.get(VT_DOMAIN_URL + value, headers=headers, timeout=30)
            resp.raise_for_status()
            j=resp.json().get('data',{}).get('attributes',{})
            stats=j.get('last_analysis_stats',{})
            malicious_count=int(stats.get('malicious',0))
            risk="High" if malicious_count >= 5 else "Suspicious" if malicious_count >= 1 else "Low"
            note=", ".join(j.get('tags',[])) if 'tags' in j else ''
            result = {'Risk Level':risk,'Malicious':malicious_count,
                        'Harmless':stats.get('harmless',0),'Suspicious':stats.get('suspicious',0),
                        'Undetected':stats.get('undetected',0),'Notes':note}
            self._update_cache(value, result)
            return result, False
        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 404:
                # Modified: Create a result for "Not Found" and cache it
                tqdm.write(f"{color.WARNING}[!] Domain {value} not found on VirusTotal. Caching this result.{color.END}")
                result = {'Risk Level': 'N/A', 'Malicious': 'N/A', 'Notes': 'Not Found on VT'}
                self._update_cache(value, result)
                return result, False
            else:
                tqdm.write(f"{color.ERROR}[✗] HTTP error for domain {value}: {e}{color.END}")
                return {'Error': str(e)}, False
        except requests.exceptions.SSLError: return self._handle_ssl_error(), False
        except RequestException as e: return {'Error':str(e)}, False
    
    def enrich_url(self, value):
        if self.ssl_error_detected: return {'Error': 'SSL verification failed; skipping further enrichment.'}
        cached_data, was_cached = self._get_from_cache(value)
        if was_cached: return cached_data, True

        self._wait_for_api('VirusTotal')
        try:
            headers = {"x-apikey": self.vt_key}
            url_id = base64.urlsafe_b64encode(value.encode()).decode().strip("=")
            resp = requests.get(VT_URLS_URL + url_id, headers=headers, timeout=30)
            resp.raise_for_status()
            j=resp.json().get('data',{}).get('attributes',{})
            stats=j.get('last_analysis_stats',{})
            malicious_count=int(stats.get('malicious',0))
            risk="High" if malicious_count>0 else "Low"
            note=", ".join(j.get('tags',[])) if 'tags' in j else ''
            result = {'Risk Level':risk,'Malicious':malicious_count,
                        'Harmless':stats.get('harmless',0),'Suspicious':stats.get('suspicious',0),
                        'Undetected':stats.get('undetected',0),'Notes':note}
            self._update_cache(value, result)
            return result, False
        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 404:
                # Modified: Create a result for "Not Found" and cache it
                tqdm.write(f"{color.WARNING}[!] URL {value} not found on VirusTotal. Caching this result.{color.END}")
                result = {'Risk Level': 'N/A', 'Malicious': 'N/A', 'Notes': 'Not Found on VT'}
                self._update_cache(value, result)
                return result, False
            else:
                tqdm.write(f"{color.ERROR}[✗] HTTP error for URL {value}: {e}{color.END}")
                return {'Error': str(e)}, False
        except requests.exceptions.SSLError: return self._handle_ssl_error(), False
        except RequestException as e: return {'Error':str(e)}, False

# === Main functions ===
def extract_iocs(file_path, scan_attachments, enable_hashing):
    """Orchestrates the IOC extraction process from files."""
    all_ips = defaultdict(lambda: {'source_files': set(), 'enrichment': {}})
    all_hashes = defaultdict(lambda: {'source_files': set(), 'enrichment': {}, 'hash_type': ''})
    all_domains = defaultdict(lambda: {'source_files': set(), 'enrichment': {}})
    all_urls = defaultdict(lambda: {'source_files': set(), 'enrichment': {}})
    all_emails = defaultdict(lambda: {'source_files': set(), 'enrichment': {}})
    all_email_data = {}
    all_generic_ips = defaultdict(lambda: {'source_files': set(), 'enrichment': {}})
    all_generic_domains = defaultdict(lambda: {'source_files': set(), 'enrichment': {}})
    all_generic_urls = defaultdict(lambda: {'source_files': set(), 'enrichment': {}})
    
    try:
        all_files = get_files_from_path(file_path)
    except ValueError as e:
        tqdm.write(f"{color.ERROR}[✗] {e}{color.END}")
        return None, None, None, None, None, None, None, None, None

    if not all_files:
        tqdm.write(f"{color.WARNING}[!] No supported files found in the specified path. Exiting.{color.END}")
        return None, None, None, None, None, None, None, None, None

    for f in tqdm(all_files, desc=f"{color.INFO}Extracting IOCs from Files{color.END}", unit="file", ncols=80, leave=False):
        tqdm.write(f"{color.INFO}[*] Processing file: {os.path.basename(f)}{color.END}")
        file_ips, file_hashes, file_domains, file_urls, file_emails_dict, file_msg_ids, extracted_pcap_files = extract_iocs_from_file(f, scan_attachments)
        
        for ioc in file_ips:
            if is_generic_ioc(ioc, "IP"):
                all_generic_ips[ioc]['source_files'].add(f)
            else:
                all_ips[ioc]['source_files'].add(f)

        for h_dict in file_hashes:
            all_hashes[h_dict['hash']]['source_files'].add(h_dict['source_file'])
            all_hashes[h_dict['hash']]['hash_type'] = h_dict.get('type', 'Unknown')
        
        for ioc in file_domains:
            if is_generic_ioc(ioc, "Domain"):
                all_generic_domains[ioc]['source_files'].add(f)
            else:
                all_domains[ioc]['source_files'].add(f)
                
        for ioc in file_urls:
            if is_generic_ioc(ioc, "URL"):
                all_generic_urls[ioc]['source_files'].add(f)
            else:
                all_urls[ioc]['source_files'].add(f)

        for ioc, msg_id in file_emails_dict.items():
            all_emails[ioc]['source_files'].add(f)
            all_email_data[ioc] = msg_id

    # The hashing logic is now correctly placed in this conditional block
    if enable_hashing:
        for f in tqdm(all_files, desc=f"{color.INFO}Hashing Source Files{color.END}", unit="file", ncols=80, leave=False):
            for h_dict in hash_file(f):
                all_hashes[h_dict['hash']]['source_files'].add(h_dict['source_file'])
                all_hashes[h_dict['hash']]['hash_type'] = h_dict.get('type', 'Unknown')

    print(f"\n{color.SUCCESS}[✓] Extraction complete.{color.END}")

    return (all_ips, all_hashes, all_domains, all_urls, all_emails, all_email_data,
            all_generic_ips, all_generic_domains, all_generic_urls)

def enrich_iocs(all_ips, all_hashes, all_domains, all_urls, all_generic_ips, all_generic_domains, all_generic_urls,
                enricher, enrich_ip_flag, enrich_hash_flag, enrich_domain_flag, enrich_url_flag, enrich_generic_flag):
    """Enriches IOCs and handles porting of contextual data like country codes."""

    # This dictionary will map a domain discovered via IP enrichment to that IP's country
    domain_to_ip_country = {}

    if enrich_ip_flag:
        newly_discovered_domains = set()
        
        # Consolidate IPs for enrichment to avoid code duplication
        ips_to_enrich = {**all_ips, **(all_generic_ips if enrich_generic_flag else {})}
        
        for ip, data in tqdm(ips_to_enrich.items(), desc=f"{color.INFO}Enriching IPs{color.END}", unit="IP", ncols=80, leave=False):
            enrichment_data, was_cached = enricher.enrich_ip(ip)
            data['enrichment'] = (enrichment_data, was_cached)
            
            if enrichment_data:
                domain_from_ip = enrichment_data.get('Domain')
                country_from_ip = enrichment_data.get('Country')
                
                if domain_from_ip:
                    # If this is a new domain, add it to the list to be enriched later
                    if domain_from_ip not in all_domains and domain_from_ip not in all_generic_domains and not enricher._get_from_cache(domain_from_ip)[0]:
                        newly_discovered_domains.add(domain_from_ip)
                    
                    # Map the discovered domain to its IP's country code
                    if country_from_ip and domain_from_ip not in domain_to_ip_country:
                        domain_to_ip_country[domain_from_ip] = country_from_ip

        if newly_discovered_domains:
            tqdm.write(f"{color.INFO}[*] {len(newly_discovered_domains)} new domains discovered from IP enrichment have been added to the queue.{color.END}")
            for new_domain in newly_discovered_domains:
                source_file_note = "Discovered via IP enrichment"
                all_domains[new_domain] = {'source_files': {source_file_note}, 'enrichment': {}}
                

    if enrich_hash_flag:
        hashes_to_enrich = {**all_hashes}
        for h, data in tqdm(hashes_to_enrich.items(), desc=f"{color.INFO}Enriching Hashes{color.END}", unit="Hash", ncols=80, leave=False):
            enrichment_data, was_cached = enricher.enrich_hash(h)
            data['enrichment'] = (enrichment_data, was_cached)

    if enrich_domain_flag:
        domains_to_enrich = {**all_domains, **(all_generic_domains if enrich_generic_flag else {})}
        for d, data in tqdm(domains_to_enrich.items(), desc=f"{color.INFO}Enriching Domains{color.END}", unit="Domain", ncols=80, leave=False):
            enrichment_data, was_cached = enricher.enrich_domain(d)
            
            if isinstance(enrichment_data, dict) and d in domain_to_ip_country and not enrichment_data.get('Country'):
                enrichment_data['Country'] = domain_to_ip_country[d]
            
            data['enrichment'] = (enrichment_data, was_cached)

    if enrich_url_flag:
        urls_to_enrich = {**all_urls, **(all_generic_urls if enrich_generic_flag else {})}
        for u, data in tqdm(urls_to_enrich.items(), desc=f"{color.INFO}Enriching URLs{color.END}", unit="URL", ncols=80, leave=False):
            enrichment_data, was_cached = enricher.enrich_url(u)
            data['enrichment'] = (enrichment_data, was_cached)

    print(f"\n{color.SUCCESS}[✓] Enrichment complete.{color.END}")
    return all_ips, all_hashes, all_domains, all_urls, all_generic_ips, all_generic_domains, all_generic_urls

# === Report Writing Functions ===

def structure_report_data(all_ips, all_hashes, all_domains, all_urls, all_emails, all_email_data, all_generic_ips, all_generic_domains, all_generic_urls, all_generic_emails, defang_flag):
    """Consolidates all IOC data into a single, structured dictionary for reporting."""

    def process_ioc_dict(ioc_dict, ioc_type):
        """Helper function to format IOC data for the report."""
        processed_list = []
        for ioc, data in ioc_dict.items():
            # Enrichment data is stored as a tuple (dictionary, was_cached). We need the dictionary.
            # Use .get() for safety in case enrichment was skipped.
            enrichment_info = data.get('enrichment', ({}, False))
            enrichment_data = enrichment_info[0] if isinstance(enrichment_info, tuple) else enrichment_info

            item = {
                'ioc': defang(ioc_type, ioc) if defang_flag else ioc,
                'source_files': list(data.get('source_files', [])),
                'enrichment': enrichment_data,
                'type': ioc_type
            }
            
            # Hashes have an extra 'hash_type' field we need to preserve
            if ioc_type == 'Hash':
                item['hash_type'] = data.get('hash_type', '')

            processed_list.append(item)
        return processed_list

    # This lambda function safely gets the integer value for sorting, defaulting to -1 if it's not an integer.
    safe_sort_key = lambda x: x['enrichment']['Malicious'] if isinstance(x['enrichment'].get('Malicious'), int) else -1

    report = {
        'iocs': {
            'ips': sorted(process_ioc_dict(all_ips, 'IP'), key=lambda x: x['enrichment'].get('Abuse Score', -1), reverse=True),
            'hashes': sorted(process_ioc_dict(all_hashes, 'Hash'), key=safe_sort_key, reverse=True),
            'domains': sorted(process_ioc_dict(all_domains, 'Domain'), key=safe_sort_key, reverse=True),
            'urls': sorted(process_ioc_dict(all_urls, 'URL'), key=safe_sort_key, reverse=True),
            'emails': process_ioc_dict(all_emails, 'Email')
        },
        'generic_iocs': {
            'ips': sorted(process_ioc_dict(all_generic_ips, 'IP'), key=lambda x: x['enrichment'].get('Abuse Score', -1), reverse=True),
            'domains': sorted(process_ioc_dict(all_generic_domains, 'Domain'), key=safe_sort_key, reverse=True),
            'urls': sorted(process_ioc_dict(all_generic_urls, 'URL'), key=safe_sort_key, reverse=True),
            'emails': process_ioc_dict(all_generic_emails, 'Email')
        }
    }
    return report

def write_json_report(report_data, timestamp):
    """Writes the report data to a JSON file."""
    output_file = f"ioc_report_{timestamp}.json"
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, indent=4)
        tqdm.write(f"{color.SUCCESS}[✓] JSON report written to {output_file}{color.END}")
    except Exception as e:
        tqdm.write(f"{color.ERROR}[✗] Failed to write JSON report: {e}{color.END}")

def write_html_report(report_data, timestamp, segregate_generic_flag):
    """Writes the report data to a visually clean and interactive HTML file."""
    output_file = f"ioc_report_{timestamp}.html"
    
    def get_risk_class(ioc_data):
        enrich = ioc_data.get('enrichment', {})
        if not enrich: return "risk-low"

        # Safely check Abuse Score
        abuse_score = enrich.get('Abuse Score')
        if isinstance(abuse_score, int):
            if abuse_score >= 70: return "risk-high"
            if abuse_score >= 30: return "risk-medium"

        # Safely check Malicious score
        malicious_score = enrich.get('Malicious')
        if isinstance(malicious_score, int) and malicious_score > 0:
            return "risk-high"

        # Check categorical Risk Level
        if 'Risk Level' in enrich:
            if enrich['Risk Level'] == "High": return "risk-high"
            if enrich['Risk Level'] == "Medium": return "risk-medium"
            
        return "risk-low"
    
    # <<<FIXED>>>: This function now handles non-integer values for 'Malicious' score.
    def get_top_iocs(report_data):
        top_iocs = []
        
        # Consolidate all IOCs from both standard and generic lists
        all_iocs_to_check = []
        for ioc_type in ['ips', 'hashes', 'domains', 'urls']:
            all_iocs_to_check.extend(report_data['iocs'].get(ioc_type, []))
            all_iocs_to_check.extend(report_data['generic_iocs'].get(ioc_type, []))
        
        for ioc in all_iocs_to_check:
            is_high_risk = False
            abuse_score = ioc['enrichment'].get('Abuse Score')
            malicious_score = ioc['enrichment'].get('Malicious')
            
            if isinstance(abuse_score, int) and abuse_score >= 30:
                is_high_risk = True
            
            if not is_high_risk and isinstance(malicious_score, int) and malicious_score > 0:
                is_high_risk = True

            if is_high_risk:
                top_iocs.append(ioc)
        
        # Sort by type first, then by score
        def sort_key(x):
            type_order = {'IP': 0, 'Domain': 1, 'URL': 2, 'Hash': 3}
            abuse_score = x['enrichment'].get('Abuse Score')
            malicious_score = x['enrichment'].get('Malicious')
            
            primary_score = 0
            if isinstance(abuse_score, int):
                primary_score = abuse_score
            elif isinstance(malicious_score, int):
                primary_score = malicious_score
                
            return (type_order.get(x['type'], 99), -primary_score) # -primary_score for descending

        return sorted(top_iocs, key=sort_key)


    html_template = f"""
    <!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><title>Grabipy IOC Report</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif; margin: 0; padding: 0; background-color: #f8f9fa; color: #212529; }}
        .container {{ max-width: 95%; margin: 20px auto; padding: 20px; background-color: white; border-radius: 8px; box-shadow: 0 4px 12px rgba(0,0,0,0.1); }}
        h1, h2 {{ color: #0056b3; border-bottom: 2px solid #dee2e6; padding-bottom: 10px; }}
        h1 {{ font-size: 2.5em; text-align: center; }} 
        h2 {{ font-size: 1.8em; margin-top: 40px; cursor: pointer; position: relative; padding: 10px; border-radius: 5px; background-color: #f2f2f2;}}
        h2::after {{ content: ' ▼'; font-size: 0.7em; position: absolute; right: 15px; top: 50%; transform: translateY(-50%); transition: transform 0.2s; }}
        h2.collapsed::after {{ transform: translateY(-50%) rotate(-90deg); }}
        table {{ width: 100%; border-collapse: collapse; margin-bottom: 30px; box-shadow: 0 4px 8px rgba(0,0,0,0.05); }}
        th, td {{ padding: 12px 15px; border: 1px solid #dee2e6; text-align: left; vertical-align: top; }}
        th {{ background-color: #007bff; color: white; font-weight: 600; }}
        tr:nth-child(even) {{ background-color: #f8f9fa; }}
        tr:hover {{ background-color: #e9ecef; }}
        .risk-high {{ border-left: 5px solid #d9534f; }}
        .risk-medium {{ border-left: 5px solid #f0ad4e; }}
        .risk-low {{ border-left: 5px solid #5cb85c; }}
        .ioc-value {{ word-break: break-all; font-family: "Courier New", Courier, monospace; font-size: 1.1em; }}
        .enrichment-key {{ font-weight: bold; color: #495057; }}
        .source-file {{ font-style: italic; color: #6c757d; font-size: 0.9em; word-break: break-all; }}
        .details-list {{ margin: 0; padding: 0; list-style-type: none; }}
        .risk-text-high {{ color: #d9534f; font-weight: bold; }}
        .risk-text-medium {{ color: #f0ad4e; font-weight: bold; }}
        .risk-text-low {{ color: #5cb85c; }}
        .nav-bar {{ position: sticky; top: 0; z-index: 1000; background-color: #343a40; padding: 15px; border-radius: 5px; margin-bottom: 30px; text-align: center; box-shadow: 0 2px 5px rgba(0,0,0,0.2); }}
        .nav-bar a {{ margin: 0 15px; color: #f8f9fa; text-decoration: none; font-weight: bold; font-size: 1.1em; }}
        .nav-bar a:hover {{ color: #007bff; }}
    </style>
    <script>
        function toggleTable(headerElement, tableId) {{
            var table = document.getElementById(tableId);
            headerElement.classList.toggle("collapsed");
            if (table.style.display === "none" || table.style.display === "") {{
                table.style.display = "table";
            }} else {{
                table.style.display = "none";
            }}
        }}
    </script>
    </head><body>
    <div class="container">
    <h1>Grabipy IOC Report</h1>
    <p style="text-align: center;">Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    
    <div class="nav-bar">
        <a href="#summary">Executive Summary</a>
        <a href="#ips">IPs</a>
        <a href="#hashes">Hashes</a>
        <a href="#domains">Domains</a>
        <a href="#emails">Emails</a>
        <a href="#generic-ips">Generic IPs</a>
        <a href="#generic-domains">Generic Domains</a>
    </div>
    """
    
    def create_table(title, iocs, table_id):
        if not iocs: return ""
        
        headers = ['IOC', 'Type', 'Risk Level', 'Score / Detections', 'Country', 'Enrichment Details', 'Source Files']
        
        table_html = f"<h2 id='{table_id}' onclick='toggleTable(this, \"{table_id}-table\")'>{title} ({len(iocs)})</h2>"
        table_html += f"<table id='{table_id}-table'>"
        table_html += "<tr>" + "".join(f"<th>{h}</th>" for h in headers) + "</tr>"
        
        for ioc_data in iocs:
            risk_class = get_risk_class(ioc_data)
            table_html += f"<tr class='{risk_class}'>"
            enrich = ioc_data.get('enrichment', {})
            
            ioc_value = ioc_data.get('ioc', '')
            ioc_type = ioc_data.get('type', '')
            if ioc_type == 'Hash':
                ioc_type += f" ({ioc_data.get('hash_type', '')})"
            table_html += f"<td class='ioc-value'>{ioc_value}</td><td>{ioc_type}</td>"

            risk_level = enrich.get('Risk Level', 'N/A')
            risk_level_class = f"risk-text-{risk_level.lower()}" if risk_level != 'N/A' else ''
            table_html += f"<td class='{risk_level_class}'>{risk_level}</td>"

            score = enrich.get('Abuse Score', enrich.get('Malicious', 'N/A'))
            table_html += f"<td>{score}</td>"

            country = enrich.get('Country', '')
            table_html += f"<td>{country}</td>"

            enrich_html = "<ul class='details-list'>"
            keys_to_exclude = {'Risk Level', 'Abuse Score', 'Malicious', 'Country'}
            
            remaining_keys = sorted([k for k in enrich.keys() if k not in keys_to_exclude])
            
            details_found = False
            for key in remaining_keys:
                value = enrich[key]
                if value:
                    enrich_html += f"<li><span class='enrichment-key'>{key}:</span> {value}</li>"
                    details_found = True

            if not details_found:
                enrich_html += "<li>No additional details</li>"

            enrich_html += "</ul>"
            table_html += f"<td>{enrich_html}</td>"

            source_files = " | ".join(ioc_data['source_files'])
            table_html += f"<td class='source-file'>{source_files}</td>"
            
            table_html += "</tr>"

        table_html += "</table>"
        return table_html

    top_iocs = get_top_iocs(report_data)
    if top_iocs:
        html_template += create_table("Executive Summary: Top IOCs for Review", top_iocs, "summary")

    main_iocs = report_data['iocs']
    generic_iocs = report_data['generic_iocs']

    if not segregate_generic_flag:
        main_iocs['ips'].extend(generic_iocs.get('ips', []))
        main_iocs['domains'].extend(generic_iocs.get('domains', []))
        main_iocs['urls'].extend(generic_iocs.get('urls', []))
        main_iocs['emails'].extend(generic_iocs.get('emails', []))
    
    html_template += create_table("IP Addresses", main_iocs['ips'], "ips")
    html_template += create_table("Hashes", main_iocs['hashes'], "hashes")
    html_template += create_table("Domains", main_iocs['domains'], "domains")
    html_template += create_table("Emails", main_iocs['emails'], "emails")

    if segregate_generic_flag:
        html_template += create_table("Generic IP Addresses", generic_iocs.get('ips', []), "generic-ips")
        html_template += create_table("Generic Domains", generic_iocs.get('domains', []), "generic-domains")
        html_template += create_table("Generic URLs", generic_iocs.get('urls', []), "generic-urls")
        html_template += create_table("Generic Emails", generic_iocs.get('emails', []), "generic-emails")

    html_template += "</div></body></html>"
    
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html_template)
        tqdm.write(f"{color.SUCCESS}[✓] HTML report written to {output_file}{color.END}")
    except Exception as e:
        tqdm.write(f"{color.ERROR}[✗] Failed to write HTML report: {e}{color.END}")

def display_guide():
    """Displays a detailed user guide and waits for user input to return."""
    os.system('cls' if os.name == 'nt' else 'clear')
    
    guide_banner = r"""
      ██████╗ ██╗   ██╗██╗██████╗ ███████╗
     ██╔════╝ ██║   ██║██║██╔══██╗██╔════╝
     ██║  ███╗██║   ██║██║██║  ██║█████╗  
     ██║   ██║██║   ██║██║██║  ██║██╔══╝  
     ╚██████╔╝╚██████╔╝██║██████╔╝███████╗
      ╚═════╝  ╚═════╝ ╚═╝╚═════╝ ╚══════╝
    """
    print(f"{color.SUCCESS}{guide_banner}{color.END}")
    
    print("-" * 60)
    print(f" {color.INFO}How to use Grabipy - Step-by-Step Guide{color.END}")
    print("-" * 60)

    print(f"\n {color.WARNING}1. First-Time Setup:{color.END}")
    print("    - When you run the script for the first time, it will check for")
    print("      all required libraries and ask for permission to install them.")
    print("    - This is a one-time process.")
    
    print(f"\n {color.WARNING}2. Setting API Keys (Option 2):{color.END}")
    print("    - For IOC enrichment, you need API keys from AbuseIPDB and VirusTotal.")
    print("    - Use Option 2 in the main menu to enter your keys. They will be")
    print("      saved securely in a 'config.ini' file for future use.")

    print(f"\n {color.WARNING}3. Running a Scan (Option 1):{color.END}")
    print("    - Select Option 1 to start.")
    print("    - You'll be asked for a file or folder path. You can provide a full")
    print("      path (e.g., 'C:\\Users\\YourUser\\Documents') or just press Enter")
    print("      to scan the current directory.")
    print("    - The script will then extract all IOCs and ask if you want to")
    print("      enrich them using your saved API keys.")

    print(f"\n {color.WARNING}4. The Output:{color.END}")
    print("    - Detailed CSV, JSON, and/or HTML reports named 'ioc_report_...' will be")
    print("      created in the script's directory. These files contain all the")
    print("      found IOCs and their enrichment data.")

    print(f"\n {color.WARNING}Pro Tip: Fast Startup:{color.END}")
    print("    - After the first run, you can start the script much faster by")
    print("      running it with the '--skip-check' flag from your terminal:")
    print(f"      {color.SUCCESS}python grabipy_v5.py --skip-check{color.END}")

    print("-" * 60)
    input(f"\n{color.INFO}Press Enter to return to the main menu...{color.END}")

def main_menu():
    """Main menu and script orchestration."""
    os.system('cls' if os.name == 'nt' else 'clear')
    
    banner = r"""
    ██████╗ ██████╗  █████╗ ██████╗ ██╗██████╗ ██╗   ██╗
    ██╔════╝ ██╔══██╗██╔══██╗██╔══██╗██║██╔══██╗╚██╗ ██╔╝
    ██║  ███╗██████╔╝███████║██████╔╝██║██████╔╝ ╚████╔╝ 
    ██║   ██║██╔══██╗██╔══██║██╔══██╗██║██╔═══╝   ╚██╔╝  
    ╚██████╔╝██║  ██║██║  ██║██████╔╝██║██║       ██║   
     ╚═════╝ ╚═╝  ╚═╝╚═╝  ╚═╝╚═════╝ ╚═╝╚═╝       ╚═╝   
    """
    
    print(f"{color.INFO}{banner}{color.END}")
    print(f"      {color.SUCCESS}Threat Intelligence & IOC Enrichment Tool{color.END}")
    print(f"            {color.WARNING}Created by Blu3J4x{color.END}\n")
    
    print("-" * 60)
    print(f" {color.INFO}Welcome to Grabipy!{color.END}")
    print(" This tool extracts and enriches Indicators of Compromise (IOCs)")
    print(" from a wide variety of file types to aid in your investigations.")
    print("-" * 60)

    while True:
        print(f"\n {color.SUCCESS}--- MAIN MENU ---{color.END}")
        print(f" {color.INFO}1.{color.END} {color.WARNING}Start Scan{color.END}      - Begin IOC extraction and enrichment.")
        print(f" {color.INFO}2.{color.END} {color.WARNING}API Keys{color.END}        - Set up or update your API keys.")
        print(f" {color.INFO}3.{color.END} {color.WARNING}Guide{color.END}           - View the user guide and instructions.")
        print(f" {color.INFO}4.{color.END} {color.WARNING}Exit{color.END}            - Close the application.")
        
        choice = input(f"\n{color.SUCCESS}Please enter your choice (1-4): {color.END}").strip()
        
        if choice == '1':
            
            all_ips, all_hashes, all_domains, all_urls, all_emails, all_email_data, all_generic_ips, all_generic_domains, all_generic_urls = [], {}, {}, {}, {}, {}, [], [], []
            all_generic_emails = defaultdict(lambda: {'source_files': set(), 'enrichment': {}})


            print("-" * 60)
            file_path = input(f" {color.INFO}Enter the path to your file or folder{color.END}\n (Default: current directory): ").strip() or "."
            scan_attachments = input(f" {color.INFO}Scan email attachments? (y/N):{color.END} ").strip().lower() in ('y','yes')
            
            # Add this new section for hashing
            while True:
                hash_input = input(f"{color.INFO}Do you wish to hash the source files? (y/n) {color.END}").lower()
                if hash_input in ['y', 'yes']:
                    enable_hashing = True
                    break
                elif hash_input in ['n', 'no']:
                    enable_hashing = False
                    break
                else:
                    print(f"{color.ERROR}[✗] Invalid input. Please enter 'y' or 'n'.{color.END}")
            
            print("-" * 60)
            
            all_ips, all_hashes, all_domains, all_urls, all_emails, all_email_data, all_generic_ips, all_generic_domains, all_generic_urls = extract_iocs(file_path, scan_attachments, enable_hashing)

            if all_ips or all_hashes or all_domains or all_urls or all_emails or all_generic_ips or all_generic_domains or all_generic_urls:
                enrich_choice = input(f"\nWould you like to proceed with enrichment? (Y/n): ").strip().lower()
                
                enricher = None
                enrich_ip_flag = False
                enrich_hash_flag = False
                enrich_domain_flag = False
                enrich_url_flag = False
                enrich_generic_flag = False

                if enrich_choice not in ('n', 'no'):
                    keys_loaded = False
                    while True:
                        abuse_key, vt_key = load_config()
                        if abuse_key and vt_key:
                            keys_loaded = True
                            enricher = Enricher(abuse_key, vt_key, load_cache())
                            break 
                        print(f"\n{color.WARNING}[!] No API keys found.{color.END}")
                        setup_now = input("Would you like to set up API keys now? (Y/n): ").strip().lower()
                        if setup_now in ('', 'y', 'yes'):
                            setup_config() 
                        else:
                            print(f"{color.INFO}[*] Skipping enrichment.{color.END}")
                            break 
                    
                    if keys_loaded:
                        enrich_ip_flag = input(f"Enrich IPs? [Y/n]: {color.END}").lower() in ('y', 'yes')
                        enrich_hash_flag = input(f"Enrich Hashes? [Y/n]: {color.END}").lower() in ('y', 'yes')
                        enrich_domain_flag = input(f"Enrich Domains? [Y/n]: {color.END}").lower() in ('y', 'yes')
                        enrich_url_flag = input(f"Enrich URLs? [Y/n]: {color.END}").lower() in ('y', 'yes')
                        enrich_generic_flag = input(f"Enrich generic IOCs (e.g., from Google, AWS)? [y/N]: {color.END}").lower() in ('y', 'yes')

                        try:
                            # The enrich_iocs function returns a set of dictionaries, so we're not unpacking a tuple here.
                            all_ips, all_hashes, all_domains, all_urls, all_generic_ips, all_generic_domains, all_generic_urls = enrich_iocs(
                                all_ips, all_hashes, all_domains, all_urls, all_generic_ips, all_generic_domains, all_generic_urls, 
                                enricher, enrich_ip_flag, enrich_hash_flag, enrich_domain_flag, enrich_url_flag, enrich_generic_flag)
                            save_cache(enricher.cache)
                            tqdm.write(f"{color.SUCCESS}[✓] Enrichment cache has been updated.{color.END}")
                        except KeyboardInterrupt:
                            tqdm.write(f"\n{color.WARNING}[!] Enrichment interrupted by user (Ctrl+C). Writing collected data to CSV...{color.END}")
                else:
                    tqdm.write(f"\n{color.INFO}[*] Skipping enrichment. Writing extracted IOCs...{color.END}")
            else:
                print(f"{color.WARNING}[!] No IOCs found. Exiting.{color.END}")
                break
            
            # --- START: Generic Provider Segregation Logic ---
            # This block runs after enrichment to move any newly identified generic IOCs
            # (e.g., an IP whose ISP is identified as 'Google LLC') to the generic lists.

            GENERIC_PROVIDERS_ISP = [
                'Google', 'Microsoft', 'Amazon', 'Cloudflare', 'Akamai', 'GoDaddy',
                'DigitalOcean', 'Comcast', 'Verizon', 'AT&T', 'Oracle',
                'IONOS', 'Aliyun', 'Palo Alto'
            ]
            
            GENERIC_DOMAINS_SET = {
                "google.com", "gmail.com", "gstatic.com", "googlevideo.com", "googleusercontent.com", "youtube.com",
                "windows.net", "microsoft.com", "msftncsi.com", "azure.com", "office.com", "outlook.com", "hotmail.com", "live.co.uk", "onedrive.live.com",
                "amazonaws.com", "s3.amazonaws.com", "cloudfront.net",
                "akamaihd.net", "apple.com", "icloud.com", "github.com", "mailchimp.com", "dropbox.com", "logmein.com", "facebook.com", "twitter.com", "linkedin.com",
                "oracle.com", "digitalocean.com", "comcast.com", "ionos.com", "alibabacloud.com", "paloaltonetworks.com"
            }

            # --- Segregate IPs based on enriched ISP data AND domain name ---
            ips_to_move = []
            for ip, data in all_ips.items():
                enrichment_info = data.get('enrichment')
                if not enrichment_info: continue

                enrichment_data = enrichment_info[0] if isinstance(enrichment_info, tuple) else enrichment_info
                if not isinstance(enrichment_data, dict): continue

                is_generic = False

                # Check 1: ISP Name
                isp = enrichment_data.get('ISP', '')
                if any(provider.lower() in isp.lower() for provider in GENERIC_PROVIDERS_ISP):
                    is_generic = True
                
                # Check 2: Enriched Domain Name
                if not is_generic:
                    enriched_domain = enrichment_data.get('Domain', '')
                    if enriched_domain and get_root_domain(enriched_domain) in GENERIC_DOMAINS_SET:
                        is_generic = True

                if is_generic:
                    ips_to_move.append(ip)

            for ip in ips_to_move:
                if ip in all_ips:
                    all_generic_ips[ip] = all_ips.pop(ip)

            # --- Segregate Domains based on the domain name ---
            domains_to_move = []
            for domain in list(all_domains.keys()): # Use list() to avoid issues with changing dict size
                if get_root_domain(domain) in GENERIC_DOMAINS_SET:
                    domains_to_move.append(domain)
                    
            for domain in domains_to_move:
                if domain in all_domains:
                    all_generic_domains[domain] = all_domains.pop(domain)

            # --- Segregate URLs based on the URL's root domain ---
            urls_to_move = []
            for url in list(all_urls.keys()):
                if get_root_domain(url) in GENERIC_DOMAINS_SET:
                    urls_to_move.append(url)

            for url in urls_to_move:
                if url in all_urls:
                    all_generic_urls[url] = all_urls.pop(url)

            # New: --- Segregate Emails based on the email's root domain ---
            emails_to_move = []
            for email in list(all_emails.keys()):
                if get_root_domain(email) in GENERIC_DOMAINS_SET:
                    emails_to_move.append(email)

            for email in emails_to_move:
                if email in all_emails:
                    all_generic_emails[email] = all_emails.pop(email)

            # --- END: Generic Provider Segregation Logic ---
            
            defang_output_flag = input("\nWould you like to defang the output (IPs, Domains, and URLs)? (y/N): ").strip().lower() in ('y','yes')
            segregate_generic_flag = input("Segregate generic IOCs into a separate section in the report? (Y/n): ").strip().lower() not in ('n', 'no')
            
            # Loop to validate output formats
            valid_formats = ['csv', 'json', 'html']
            while True:
                output_formats_input = input("\nEnter desired output formats (csv, json, html) - separate with commas: ").strip().lower()
                if not output_formats_input:
                    requested_formats = ['csv']
                    print(f"{color.INFO}[*] No format selected, defaulting to CSV.{color.END}")
                    break
                
                requested_formats = [f.strip() for f in output_formats_input.split(',')]
                
                if all(fmt in valid_formats for fmt in requested_formats):
                    break
                else:
                    print(f"{color.ERROR}[✗] Invalid format. Please enter a comma-separated list of 'csv', 'json', and/or 'html'.{color.END}")

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Pass all necessary flags to structure_report_data
            report_data = structure_report_data(
                all_ips, all_hashes, all_domains, all_urls, all_emails, all_email_data,
                all_generic_ips, all_generic_domains, all_generic_urls, all_generic_emails,
                defang_output_flag
            )

            # --- Write reports based on user selection ---
            if 'csv' in requested_formats:
                try:
                    output_file_csv = f"ioc_report_{timestamp}.csv"
                    
                    # --- This is the original, fixed CSV writing logic ---
                    with open(output_file_csv, 'w', newline='', encoding='utf-8') as f:
                        writer = csv.writer(f)
                        
                        main_iocs = report_data['iocs']
                        generic_iocs = report_data['generic_iocs']

                        if not segregate_generic_flag:
                            main_iocs['ips'].extend(generic_iocs.get('ips', []))
                            main_iocs['domains'].extend(generic_iocs.get('domains', []))
                            main_iocs['urls'].extend(generic_iocs.get('urls', []))
                            main_iocs['emails'].extend(generic_iocs.get('emails', []))
                            generic_iocs = {}

                        def write_csv_section(title, ioc_list, ioc_type):
                            if not ioc_list: return
                            writer.writerow([''])
                            writer.writerow([title])
                            
                            base_headers = ['IOC', 'Type']
                            enrich_headers_set = set()
                            for ioc in ioc_list:
                                enrich_headers_set.update(ioc.get('enrichment', {}).keys())
                            
                            enrich_headers = sorted(list(enrich_headers_set))

                            if ioc_type == 'Hash': base_headers.insert(1, 'Hash Type')
                            final_headers = base_headers + enrich_headers + ['Source File(s)']
                            writer.writerow(final_headers)

                            for ioc in ioc_list:
                                enrich_data = ioc.get('enrichment', {})
                                row = [
                                    ioc.get('ioc', ''),
                                    ioc.get('type', '')
                                ]
                                if ioc_type == 'Hash': row.insert(1, ioc.get('hash_type', ''))

                                for header in enrich_headers:
                                    row.append(enrich_data.get(header, ''))
                                
                                row.append(" | ".join(ioc.get('source_files', [])))
                                writer.writerow(row)

                        write_csv_section('IP IOCs', main_iocs['ips'], 'IP')
                        write_csv_section('Hash IOCs', main_iocs['hashes'], 'Hash')
                        write_csv_section('Domain IOCs', main_iocs['domains'], 'Domain')
                        write_csv_section('URL IOCs', main_iocs['urls'], 'URL')
                        write_csv_section('Email IOCs', main_iocs['emails'], 'Email')

                        if segregate_generic_flag and generic_iocs:
                            write_csv_section('Generic IP IOCs', generic_iocs.get('ips', []), 'IP')
                            write_csv_section('Generic Domain IOCs', generic_iocs.get('domains', []), 'Domain')
                            write_csv_section('Generic URL IOCs', generic_iocs.get('urls', []), 'URL')
                            write_csv_section('Generic Email IOCs', generic_iocs.get('emails', []), 'Email')
                        
                    tqdm.write(f"\n{color.SUCCESS}[✓] CSV report written to {output_file_csv}{color.END}")
                except Exception as e:
                    tqdm.write(f"{color.ERROR}[✗] An error occurred while writing the CSV file: {e}{color.END}")

            if 'json' in requested_formats:
                write_json_report(report_data, timestamp)
            
            if 'html' in requested_formats:
                write_html_report(report_data, timestamp, segregate_generic_flag)

            # --- Summary ---
            ioc_counts = {k: len(v) for k, v in report_data['iocs'].items() if v}
            generic_counts = {f"Generic {k.capitalize()}": len(v) for k, v in report_data['generic_iocs'].items() if v}
            summary = {**ioc_counts, **generic_counts}

            tqdm.write("\n=== Summary ===")
            for k, v in summary.items():
                tqdm.write(f"{k.replace('ips', 'IPs').replace('urls', 'URLs')}: {v}")
            
            print(f"\n{color.SUCCESS}[✓] All checks completed.{color.END}")
            break

        elif choice == '2':
            abuse_key, vt_key = setup_config()
            os.system('cls' if os.name == 'nt' else 'clear')
            main_menu()
            break
            
        elif choice == '3':
            display_guide()
            os.system('cls' if os.name == 'nt' else 'clear')
            main_menu()
            break

        elif choice == '4':
            print(f"{color.INFO}Exiting.{color.END}")
            sys.exit(0)
            
        else:
            print(f"{color.ERROR}[✗] Invalid choice. Please enter 1, 2, 3, or 4.{color.END}")
if __name__=="__main__":
    if '--skip-check' not in sys.argv:
        check_and_install_dependencies()
    else:
        print(f"{color.INFO}[*] Skipping dependency check as requested.{color.END}")
    
    main_menu()

